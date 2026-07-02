from app.core.parser import DocumentParser, ParsedDocument


class DocxParser(DocumentParser):
    @property
    def extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, path: str, source: str) -> ParsedDocument:
        import docx

        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return ParsedDocument(text="\n".join(parts), source=source)
