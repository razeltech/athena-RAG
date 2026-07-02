import csv

import pytest

from app.adapters.parsers.csv_parser import CsvParser
from app.adapters.parsers.docx_parser import DocxParser
from app.adapters.parsers.pdf_parser import PdfParser
from app.adapters.parsers.registry import ParserRegistry
from app.adapters.parsers.xlsx_parser import XlsxParser


def test_docx_parser_extracts_paragraphs_and_tables(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("Hello Athena")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "key"
    table.rows[0].cells[1].text = "value"
    path = tmp_path / "sample.docx"
    doc.save(path)

    parsed = DocxParser().extract(str(path), "sample.docx")
    assert "Hello Athena" in parsed.text
    assert "key" in parsed.text and "value" in parsed.text


def test_xlsx_parser_extracts_rows(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["name", "amount"])
    ws.append(["widget", 42])
    path = tmp_path / "sample.xlsx"
    wb.save(path)

    parsed = XlsxParser().extract(str(path), "sample.xlsx")
    assert "Sheet1" in parsed.text
    assert "widget" in parsed.text and "42" in parsed.text


def test_csv_parser_extracts_rows(tmp_path):
    path = tmp_path / "sample.csv"
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["name", "amount"])
        writer.writerow(["widget", "42"])

    parsed = CsvParser().extract(str(path), "sample.csv")
    assert "widget" in parsed.text and "42" in parsed.text


def _build_pdf_bytes(text: str | None) -> bytes:
    """Hand-build a minimal single-page PDF (with correct xref offsets) so
    tests don't need a real PDF fixture file or an extra dependency."""
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 200 200] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream = f"BT /F1 12 Tf 10 100 Td ({text}) Tj ET".encode() if text else b""
    objs.append(b"<< /Length %d >>\nstream\n" % len(stream) + stream + b"\nendstream")

    out = b"%PDF-1.4\n"
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(out)
    n = len(objs) + 1
    out += f"xref\n0 {n}\n".encode() + b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {n} /Root 1 0 R >>\nstartxref\n{xref_offset}\n".encode()
    out += b"%%EOF\n"
    return out


def test_pdf_parser_extracts_text(tmp_path):
    path = tmp_path / "sample.pdf"
    path.write_bytes(_build_pdf_bytes("Hello Athena"))

    parsed = PdfParser().extract(str(path), "sample.pdf")
    assert "Hello Athena" in parsed.text


def test_pdf_parser_raises_on_no_extractable_text(tmp_path):
    path = tmp_path / "blank.pdf"
    path.write_bytes(_build_pdf_bytes(None))

    with pytest.raises(ValueError):
        PdfParser().extract(str(path), "blank.pdf")


def test_registry_supports_new_formats():
    supported = ParserRegistry().supported()
    for ext in [".docx", ".xlsx", ".csv", ".pdf"]:
        assert ext in supported
